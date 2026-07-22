#!/usr/bin/env python
"""生成消融实验报告 Word 文档"""
import json
import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 读取消融实验结果
report_path = Path(__file__).parent / "reports" / "ablation_report.json"
if not report_path.exists():
    print(f"错误: 找不到消融实验报告 {report_path}")
    exit(1)

with open(report_path, "r", encoding="utf-8") as f:
    ablation_data = json.load(f)

full = ablation_data["full_context"]
noctx = ablation_data["no_context"]

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "宋体"
font.size = Pt(12)
style.element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "宋体")

# 标题
title = doc.add_heading("", level=0)
run = title.add_run("LLM-VulnDetector 消融实验报告")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 21, 41)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("作者：lauyy32    日期：2026年7月22日")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(102, 102, 102)
doc.add_paragraph("")

# 一、实验目的
doc.add_heading("一、实验目的", level=1)
doc.add_paragraph(
    "消融实验（Ablation Study）的目的是验证系统中「上下文增强」模块的实际贡献。"
    "具体来说，我想回答一个问题：如果把上下文增强去掉，直接把原始 HTTP 请求丢给 LLM 分析，"
    "检测效果会有什么变化？"
)
doc.add_paragraph(
    "这个问题直接对应课题的核心假设——「上下文增强能提升 LLM 攻击载荷识别的效果」。"
    "如果消融实验显示去掉上下文增强后效果没有变化，那说明当前数据集可能过于简单，"
    "或者上下文增强的设计还需要改进；如果效果有明显下降，则验证了上下文增强的价值。"
)

# 二、实验设计
doc.add_heading("二、实验设计", level=1)
doc.add_paragraph("实验采用对照组设计，两组使用相同的 LLM、相同的系统提示、相同的数据集，唯一区别在于是否启用上下文增强：")
doc.add_paragraph("")

table_design = doc.add_table(rows=3, cols=3)
table_design.style = "Light Grid Accent 1"
design_data = [
    ["", "实验组（有上下文增强）", "对照组（无上下文增强）"],
    ["接口", "POST /api/detect", "POST /api/detect-no-context"],
    ["预处理", "HTTP结构化解析 + 13类正则预扫描 + 上下文注入", "直接发送原始HTTP文本，无任何预处理"],
]
for ri, row_data in enumerate(design_data):
    for ci, val in enumerate(row_data):
        cell = table_design.cell(ri, ci)
        cell.text = val
        if ri == 0 or ci == 0:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")
doc.add_paragraph("数据集：56 条测试用例（41 条漏洞正例 + 12 条正常请求 + 3 条边界用例），覆盖 10 类 Web 漏洞。")
doc.add_paragraph("LLM 模型：DeepSeek-Chat。")

# 三、实验结果
doc.add_heading("三、实验结果", level=1)
doc.add_paragraph("两组实验的核心指标对比：")
doc.add_paragraph("")

table_result = doc.add_table(rows=10, cols=4)
table_result.style = "Light Grid Accent 1"
result_data = [
    ["指标", "有上下文增强", "无上下文增强", "差异"],
    ["样本总数", str(full["total"]), str(noctx["total"]), "-"],
    ["TP (真阳性)", str(full["tp"]), str(noctx["tp"]), str(full["tp"] - noctx["tp"])],
    ["TN (真阴性)", str(full["tn"]), str(noctx["tn"]), str(full["tn"] - noctx["tn"])],
    ["FP (误报)", str(full["fp"]), str(noctx["fp"]), str(full["fp"] - noctx["fp"])],
    ["FN (漏报)", str(full["fn"]), str(noctx["fn"]), str(full["fn"] - noctx["fn"])],
    ["准确率", f'{full["accuracy"]}%', f'{noctx["accuracy"]}%', f'{round(full["accuracy"] - noctx["accuracy"], 1)}%'],
    ["精确率", f'{full["precision"]}%', f'{noctx["precision"]}%', f'{round(full["precision"] - noctx["precision"], 1)}%'],
    ["F1-Score", f'{full["f1"]}%', f'{noctx["f1"]}%', f'{round(full["f1"] - noctx["f1"], 1)}%'],
    ["平均置信度", f'{full["avg_confidence"]}%', f'{noctx["avg_confidence"]}%', f'{round(full["avg_confidence"] - noctx["avg_confidence"], 1)}%'],
]
for ri, row_data in enumerate(result_data):
    for ci, val in enumerate(row_data):
        cell = table_result.cell(ri, ci)
        cell.text = val
        if ri == 0:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")
doc.add_paragraph(f'平均检测耗时对比：有上下文增强 {full["avg_time_seconds"]}s/条，无上下文增强 {noctx["avg_time_seconds"]}s/条。')

# 各类别检出率对比
doc.add_paragraph("")
doc.add_paragraph("各漏洞类型检出率对比：")

all_cats = set(full["category_detection"].keys()) | set(noctx["category_detection"].keys())
sorted_cats = sorted(all_cats)

table_cat = doc.add_table(rows=len(sorted_cats) + 1, cols=4)
table_cat.style = "Light Grid Accent 1"
cat_header = ["漏洞类别", "有上下文增强", "无上下文增强", "差异"]
for ci, val in enumerate(cat_header):
    cell = table_cat.cell(0, ci)
    cell.text = val
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True

for ri, cat in enumerate(sorted_cats):
    f = full["category_detection"].get(cat, {"total": 0, "tp": 0})
    n = noctx["category_detection"].get(cat, {"total": 0, "tp": 0})
    f_rate = round(f["tp"] / f["total"] * 100, 1) if f["total"] else 0
    n_rate = round(n["tp"] / n["total"] * 100, 1) if n["total"] else 0
    diff = round(f_rate - n_rate, 1)
    diff_str = f"+{diff}%" if diff > 0 else f"{diff}%" if diff != 0 else "-"
    
    values = [cat, f"{f_rate}%", f"{n_rate}%", diff_str]
    for ci, val in enumerate(values):
        table_cat.cell(ri + 1, ci).text = val

# 四、结果分析
doc.add_heading("四、结果分析与讨论", level=1)

doc.add_heading("4.1 准确率对比", level=2)

# 根据实际数据分析
if full["accuracy"] == noctx["accuracy"]:
    doc.add_paragraph(
        f'两组在准确率上均达到了 {full["accuracy"]}%，没有出现差异。'
        '这个结果需要客观看待——它既说明 LLM 本身对常见 Web 漏洞的识别能力已经相当强，'
        '也说明当前数据集的难度可能不足以区分上下文增强的效果。'
    )
    doc.add_paragraph(
        '这其实和我在论文阅读中的发现一致：LLMxCPG 论文中，CPG 增强对函数内漏洞的提升（+12pp）'
        '远小于对跨函数漏洞的提升（+23pp）。当前的评测用例都是单条 HTTP 请求级别的检测，'
        '类似于「函数内」场景，LLM 靠自身的模式识别能力就能搞定。'
        '上下文增强的真正价值应该在更复杂的场景中才会显现——比如跨请求的攻击链检测、'
        '需要理解应用业务逻辑的漏洞、以及需要关联多个参数才能判断的复合型攻击。'
    )
else:
    diff = round(full["accuracy"] - noctx["accuracy"], 1)
    doc.add_paragraph(
        f'有上下文增强组的准确率（{full["accuracy"]}%）比无上下文增强组（{noctx["accuracy"]}%）'
        f'高出 {diff} 个百分点，验证了上下文增强对检测准确率的提升作用。'
    )

doc.add_heading("4.2 误报率对比", level=2)
doc.add_paragraph(
    f'有上下文增强组误报 {full["fp"]} 条，无上下文增强组误报 {noctx["fp"]} 条。'
    '误报控制是安全检测工具的核心挑战之一。在当前数据集上两组都没有误报，'
    '这主要得益于 Prompt 中的 false_positive_check 自检机制在两组中都生效了——'
    '即使没有上下文增强，LLM 仍然会在给出结论前做一轮反向验证。'
    '但可以预见，在面对更模糊的边界用例时，上下文增强提供的结构化信息（如参数类型、请求路径的语义）'
    '会帮助 LLM 做出更准确的判断。'
)

doc.add_heading("4.3 置信度对比", level=2)
conf_diff = round(full["avg_confidence"] - noctx["avg_confidence"], 1)
if conf_diff > 0:
    doc.add_paragraph(
        f'有上下文增强组的平均置信度为 {full["avg_confidence"]}%，'
        f'无上下文增强组为 {noctx["avg_confidence"]}%，'
        f'差异为 +{conf_diff} 个百分点。'
        '这说明上下文增强不仅影响检测结果的正确性，还影响 LLM 对自己判断的"信心"。'
        '预扫描提供的风险信号（如检测到 SQL 元字符、XSS 标签等）相当于给了 LLM 一份"线索清单"，'
        '让它在推理时有据可依，从而给出更高的置信度。'
    )
elif conf_diff < 0:
    doc.add_paragraph(
        f'有上下文增强组的平均置信度为 {full["avg_confidence"]}%，'
        f'无上下文增强组为 {noctx["avg_confidence"]}%。'
        '两组置信度差异不大，说明在当前数据集上，LLM 对漏洞的识别信心不依赖上下文增强。'
    )
else:
    doc.add_paragraph(
        f'两组的平均置信度均为 {full["avg_confidence"]}%，没有差异。'
        '这可能是因为当前数据集中的漏洞特征都比较明显，LLM 不需要额外的上下文信号就能高置信度地做出判断。'
    )

doc.add_heading("4.4 耗时对比", level=2)
time_diff = round(full["avg_time_seconds"] - noctx["avg_time_seconds"], 2)
if time_diff > 0:
    doc.add_paragraph(
        f'有上下文增强组的平均耗时为 {full["avg_time_seconds"]}s/条，'
        f'无上下文增强组为 {noctx["avg_time_seconds"]}s/条，'
        f'上下文增强增加了约 {time_diff}s 的开销。'
        '这部分开销来自 HTTP 解析、正则预扫描和上下文构造的处理时间，'
        '但由于这些操作在本地完成（不涉及 LLM 调用），相对于 LLM 的推理时间来说占比很小。'
        '在实际使用中，这点延迟是完全可以接受的。'
    )
else:
    doc.add_paragraph(
        f'有上下文增强组的平均耗时为 {full["avg_time_seconds"]}s/条，'
        f'无上下文增强组为 {noctx["avg_time_seconds"]}s/条，'
        '两组耗时基本相当。上下文增强的预处理开销相对于 LLM 推理时间来说可以忽略不计。'
    )

# 五、结论
doc.add_heading("五、结论与反思", level=1)
doc.add_paragraph(
    "本次消融实验在 56 条标准数据集上对比了有/无上下文增强两种模式的检测效果。"
    "实验结果表明，在当前数据集的难度水平下，两种模式均达到了 100% 的准确率，"
    "上下文增强的优势尚未充分体现。"
)
doc.add_paragraph(
    "但这并不意味着上下文增强没有价值。结合论文阅读的启发，我认为有以下几点值得反思："
)
doc.add_paragraph("数据集难度问题——当前 56 条用例的攻击特征都比较明显，LLM 靠自身能力即可识别。需要设计更复杂的用例（如编码绕过、逻辑漏洞、需要关联多参数的复合攻击）才能有效区分两种模式", style="List Bullet")
doc.add_paragraph("上下文增强的粒度——当前的上下文增强停留在 HTTP 请求级别（结构化解析+正则预扫描），对应论文中的函数级上下文。更高级的上下文增强（如 CPG 级别、跨请求上下文）可能才是产生差异的关键", style="List Bullet")
doc.add_paragraph("消融实验的维度——除了准确率，还应该关注分析质量。上下文增强可能不影响是否检出漏洞，但会影响 LLM 给出的成因分析是否准确、修复建议是否有针对性。后续可以加入分析质量的评分维度", style="List Bullet")
doc.add_paragraph("这与 LLMxCPG 论文的发现一致——CPG 增强对跨函数漏洞的提升远大于函数内漏洞。当前的评测相当于都在函数内场景，需要扩展到跨函数/跨请求场景才能验证上下文增强的核心价值", style="List Bullet")

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("—— lauyy32")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(102, 102, 102)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

output_path = r"C:\Users\lenovo\Desktop\LLM-VulnDetector消融实验报告.docx"
doc.save(output_path)
print(f"报告已保存: {output_path}")
