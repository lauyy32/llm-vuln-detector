#!/usr/bin/env python
"""生成评测报告 Word 文档"""
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "宋体"
font.size = Pt(12)
style.element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "宋体")

# 标题
title = doc.add_heading("", level=0)
run = title.add_run("LLM-VulnDetector 漏洞检测系统评测报告")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 21, 41)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("作者：lauyy32    日期：2026年7月22日")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(102, 102, 102)
doc.add_paragraph("")

# 一、实验目的
doc.add_heading("一、实验目的", level=1)
doc.add_paragraph(
    "为了验证 LLM-VulnDetector 系统在 HTTP 漏洞检测任务上的实际效果，"
    "我构建了一套覆盖 10 类常见 Web 漏洞的标准评测数据集，"
    "对系统进行了全量测试。本次评测的核心目标是量化回答以下问题："
)
doc.add_paragraph("系统能否准确识别各类 Web 漏洞？（召回率）", style="List Bullet")
doc.add_paragraph("系统是否会将正常请求误判为漏洞？（误报率）", style="List Bullet")
doc.add_paragraph("系统给出的漏洞判定是否可信？（精确率与置信度）", style="List Bullet")
doc.add_paragraph("系统的响应速度能否满足实际使用需求？（平均耗时）", style="List Bullet")

# 二、实验环境
doc.add_heading("二、实验环境与设置", level=1)
doc.add_paragraph("后端框架：Python FastAPI + httpx 异步 HTTP 客户端", style="List Bullet")
doc.add_paragraph("大语言模型：DeepSeek-Chat（DeepSeek API）", style="List Bullet")
doc.add_paragraph("前端框架：Vue 3 + Element Plus", style="List Bullet")
doc.add_paragraph("检测流程：HTTP 请求解析 -> 正则预扫描 -> 上下文结构化构造 -> LLM 分析 -> JSON 结果输出", style="List Bullet")
doc.add_paragraph("Prompt 策略：系统提示 + few-shot 示例 + CoT 推理 + 降误报自检机制", style="List Bullet")
doc.add_paragraph("数据集规模：56 条测试用例（41 条漏洞正例 + 12 条正常请求 + 3 条边界用例）", style="List Bullet")

# 三、评测数据集
doc.add_heading("三、评测数据集说明", level=1)
doc.add_paragraph(
    "我手工构造了一份包含 56 条测试用例的评测数据集，覆盖 OWASP Top 10 中的 10 类漏洞。"
    "每条用例包含一段完整的原始 HTTP 请求文本，以及标注信息（是否含漏洞、漏洞类型、最低期望置信度）。"
    "数据集的分布如下表所示："
)

table = doc.add_table(rows=13, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["漏洞类别", "正例数量", "负例数量", "说明"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True

data = [
    ["SQL 注入", "6", "-", "含布尔盲注、联合查询、时间盲注等变体"],
    ["XSS（跨站脚本）", "6", "-", "反射型、存储型，含编码绕过变体"],
    ["命令注入", "5", "-", "含管道符、反引号、$() 等变体"],
    ["路径穿越", "4", "-", "含 ../ 序列、编码绕过"],
    ["SSRF", "4", "-", "含内网地址、协议转换"],
    ["XXE", "3", "-", "含外部实体注入、参数实体"],
    ["SSTI", "4", "-", "含 Jinja2、Twig 模板注入"],
    ["NoSQL 注入", "3", "-", "含 MongoDB 操作符注入"],
    ["开放重定向", "3", "-", "含任意 URL 跳转"],
    ["文件上传", "3", "-", "含恶意扩展名、双扩展名"],
    ["正常请求", "-", "12", "合理的业务请求，不含攻击载荷"],
    ["边界用例", "-", "3", "含特殊字符但非攻击的请求"],
]
for ri, row_data in enumerate(data):
    for ci, val in enumerate(row_data):
        table.cell(ri + 1, ci).text = val

doc.add_paragraph("")
doc.add_paragraph(
    "其中正常请求和边界用例作为负例，用于检验系统的误报控制能力。"
    "边界用例包含 URL 中带有特殊字符（如单引号、分号）但不构成攻击的请求，"
    "这类用例是检验系统会不会草木皆兵的关键。"
)

# 四、评测结果
doc.add_heading("四、评测结果", level=1)
doc.add_paragraph("全量 56 条测试用例的评测结果如下：")

table2 = doc.add_table(rows=8, cols=2)
table2.style = "Light List Accent 1"
metrics = [
    ["指标", "数值"],
    ["样本总数", "56"],
    ["准确率 (Accuracy)", "100.0%"],
    ["精确率 (Precision)", "100.0%"],
    ["召回率 (Recall)", "100.0%"],
    ["F1-Score", "100.0%"],
    ["误报率 (FPR)", "0.0%"],
    ["漏报率 (FNR)", "0.0%"],
]
for ri, row_data in enumerate(metrics):
    for ci, val in enumerate(row_data):
        cell = table2.cell(ri, ci)
        cell.text = val
        if ri == 0:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")
doc.add_paragraph("混淆矩阵：")

table3 = doc.add_table(rows=3, cols=3)
table3.style = "Light Grid Accent 1"
cm = [
    ["", "预测为漏洞", "预测为安全"],
    ["实际为漏洞", "TP = 41", "FN = 0"],
    ["实际为安全", "FP = 0", "TN = 15"],
]
for ri, row_data in enumerate(cm):
    for ci, val in enumerate(row_data):
        cell = table3.cell(ri, ci)
        cell.text = val
        if ri == 0 or ci == 0:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")
doc.add_paragraph("各类别检测详情：")

table4 = doc.add_table(rows=13, cols=4)
table4.style = "Light Grid Accent 1"
cat_data = [
    ["漏洞类别", "用例数", "正确检出", "平均置信度"],
    ["SQL 注入", "6", "6", "92.8%"],
    ["XSS", "6", "6", "91.7%"],
    ["命令注入", "5", "5", "95.0%"],
    ["路径穿越", "4", "4", "95.0%"],
    ["SSRF", "4", "4", "95.0%"],
    ["XXE", "3", "3", "95.0%"],
    ["SSTI", "4", "4", "91.3%"],
    ["NoSQL 注入", "3", "3", "91.7%"],
    ["开放重定向", "3", "3", "91.7%"],
    ["文件上传", "3", "3", "95.0%"],
    ["正常请求", "12", "12", "-"],
    ["边界用例", "3", "3", "-"],
]
for ri, row_data in enumerate(cat_data):
    for ci, val in enumerate(row_data):
        cell = table4.cell(ri, ci)
        cell.text = val
        if ri == 0:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph("")
doc.add_paragraph(
    "平均检测耗时：2.06 秒/条（漏洞请求平均 2.32s，正常请求平均 1.16s，"
    "LLM 对正常请求的判定更快，因为不需要生成详细的漏洞分析）。"
    "平均置信度：93.3%（仅统计真阳性样本）。"
)

# 五、结果分析
doc.add_heading("五、结果分析与讨论", level=1)

doc.add_heading("5.1 整体表现", level=2)
doc.add_paragraph(
    "56 条用例全部判定正确，零误报、零漏报。这个结果说明在当前数据集覆盖范围内，"
    "系统的上下文增强机制（HTTP 结构化解析 + 13 类正则预扫描 + 上下文注入）"
    "能够有效辅助 LLM 做出准确判断。但我需要强调，这只是一个初步验证，"
    "56 条用例的规模相对有限，真实场景中的攻击变体远比这复杂。"
)

doc.add_heading("5.2 误报控制", level=2)
doc.add_paragraph(
    "12 条正常请求和 3 条边界用例全部正确判定为安全，没有出现草木皆兵的情况。"
    "这一点我比较在意，因为我之前在实际使用一些 SAST 工具时经常被误报淹没。"
    "系统能区分边界用例（如 URL 中含单引号但不构成注入的搜索请求）和真正的攻击，"
    "我认为这主要得益于两个设计：一是 Prompt 中的 false_positive_check 自检机制，"
    "要求 LLM 在给出结论前先做一轮反向验证；二是上下文增强提供的结构化信息，"
    "让 LLM 能看到完整的请求上下文而非仅看参数值。"
)

doc.add_heading("5.3 置信度分布", level=2)
doc.add_paragraph(
    "41 个真阳性样本的平均置信度为 93.3%，最低 85%，最高 95%。"
    "没有出现勉强检出的情况（即置信度低于 80% 的判定），"
    "说明系统在检出漏洞时整体是比较有把握的。"
    "不同漏洞类型的置信度略有差异——命令注入、路径穿越、SSRF、XXE、文件上传这几类的置信度普遍在 95%，"
    "而 SSTI 和 NoSQL 注入的个别用例置信度稍低（85%），"
    "这可能跟这两类漏洞的 payload 特征不够鲜明有关，后续可以考虑针对性优化。"
)

doc.add_heading("5.4 性能表现", level=2)
doc.add_paragraph(
    "平均 2.06 秒/条的检测速度对于交互式使用场景是可以接受的。"
    "正常请求的判定明显更快（平均 1.16s），因为 LLM 只需要输出未发现漏洞的简短结论；"
    "而漏洞请求需要生成详细的成因分析、Payload 定位和修复建议，输出更长，耗时相应增加。"
    "如果后续需要做大规模批量扫描，可以考虑引入并发控制和结果缓存。"
)

# 六、结论
doc.add_heading("六、结论与后续方向", level=1)
doc.add_paragraph(
    "本次评测验证了 LLM-VulnDetector 在覆盖 10 类 Web 漏洞的标准数据集上达到了 100% 的准确率，"
    "零误报、零漏报。系统的上下文增强机制和降误报自检策略在当前测试范围内是有效的。"
)
doc.add_paragraph("后续需要推进的方向：")
doc.add_paragraph("扩大数据集规模——56 条只是起点，后续需要引入真实漏洞库（如 DVWA 靶场）的测试样本", style="List Bullet")
doc.add_paragraph("增加对抗性测试——使用编码绕过、WAF 绕过等高级攻击变体检验系统的鲁棒性", style="List Bullet")
doc.add_paragraph("与传统工具对比——将系统与 SQLMap、Burp Suite Active Scan 等成熟工具进行横向对比", style="List Bullet")
doc.add_paragraph("消融实验——量化上下文增强各组成部分的贡献，这部分已在另一份报告中详述", style="List Bullet")

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("—— lauyy32")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(102, 102, 102)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

output_path = r"C:\Users\lenovo\Desktop\LLM-VulnDetector评测报告.docx"
doc.save(output_path)
print(f"报告已保存: {output_path}")
